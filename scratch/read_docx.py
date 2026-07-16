import docx
import os

doc_path = r"C:\Users\DELL\Downloads\[QTKDS] MẪU BÁO CÁO TỔNG HỢP KẾT QUẢ CÔNG VIỆC HÀNG NGÀY (WORKLANE).docx"
output_path = r"c:\Users\DELL\Desktop\AI-Agent\AI_PhantichchisoDT\scratch\docx_content_detailed.txt"

def read_docx(path):
    doc = docx.Document(path)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("=== PARAGRAPHS ===\n")
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"[{idx}]: {para.text}\n")
                
        f.write("\n=== TABLES ===\n")
        for t_idx, table in enumerate(doc.tables):
            f.write(f"\nTable {t_idx} (Rows: {len(table.rows)}, Cols: {len(table.columns)}):\n")
            for r_idx, row in enumerate(table.rows):
                f.write(f"  Row {r_idx}:\n")
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().replace("\n", " ")
                    f.write(f"    Col {c_idx}: {cell_text}\n")

if __name__ == "__main__":
    if os.path.exists(doc_path):
        read_docx(doc_path)
        print("Done writing to", output_path)
    else:
        print("File does not exist")
